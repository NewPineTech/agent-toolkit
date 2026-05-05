"""
step3_form_cards.py
===================
Generate Markdown "form card" cho mỗi biểu mẫu trong inventory.

Input: data/inventory.csv (rows where target_kb == "forms_kb")
Output:
    - data/forms_md/{form_code}.md (1 file Markdown / biểu mẫu)
    - data/forms_generation_log.csv (tracking status)
    - data/forms_kb_rejected.csv (file thiếu metadata bắt buộc)

Pipeline cho mỗi biểu mẫu:
    1. Validate 3 field bắt buộc (form_code, download_url, department) - nếu thiếu → reject
    2. Download file (.doc/.xls/.docx/.xlsx) → extract structure (headers/columns)
    3. Lookup process_name từ process_code (cross-ref với SOP files đã OCR)
    4. Gọi LLM (gpt-4o-mini) sinh "Mục đích sử dụng" + "Từ khóa"
    5. Render Markdown card theo template
    6. Lưu file

Usage:
    python scripts/step3_form_cards.py             # full
    python scripts/step3_form_cards.py --limit 10  # test
    python scripts/step3_form_cards.py --resume    # skip files đã có .md
"""

import argparse
import csv
import io
import re
import subprocess
import sys
import time
from pathlib import Path

from tenacity import retry, stop_after_attempt, wait_exponential
from tqdm import tqdm

sys.path.insert(0, str(Path(__file__).parent))
from common import (
    PROJECT_ROOT,
    gdrive_download_bytes,
    get_gdrive_client,
    load_config,
    safe_filename,
    setup_logger,
)

logger = setup_logger("step3_form_cards", "step3_form_cards.log")


# ============================================================
# FILE STRUCTURE EXTRACTION
# ============================================================

def extract_doc_structure(file_bytes: bytes, ext: str) -> str:
    """
    Extract structure (headers, table columns, field labels) từ file biểu mẫu.
    KHÔNG extract data trống của template - chỉ lấy schema.
    Returns: text mô tả structure (max ~500 chars).
    """
    try:
        if ext == "docx":
            return _extract_docx_structure(file_bytes)
        elif ext == "doc":
            # legacy .doc - khó parse trong Python pure, fallback dùng tên file
            return "(File định dạng .doc cũ - không parse được structure)"
        elif ext == "xlsx":
            return _extract_xlsx_structure(file_bytes)
        elif ext == "xls":
            return _extract_xls_structure(file_bytes)
        else:
            return f"(Định dạng {ext} không hỗ trợ extract)"
    except Exception as e:
        logger.debug(f"Structure extraction failed: {e}")
        return "(Không extract được structure)"


def _extract_docx_structure(file_bytes: bytes) -> str:
    """Extract heading + table column names from .docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # Headings
    for para in doc.paragraphs[:30]:
        if para.style.name.startswith("Heading") and para.text.strip():
            parts.append(f"Heading: {para.text.strip()}")
        elif para.text.strip() and len(para.text.strip()) < 100 and para.text.strip().endswith(":"):
            parts.append(f"Field: {para.text.strip()}")

    # Table column names (only header row of each table)
    for table in doc.tables[:5]:
        if table.rows:
            cols = [cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()]
            if cols:
                parts.append(f"Table columns: {' | '.join(cols[:8])}")

    result = "\n".join(parts[:15])
    return result[:500] if result else "(Template trống)"


def _extract_xlsx_structure(file_bytes: bytes) -> str:
    """Extract sheet names + first row (headers) từ .xlsx."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames[:3]:
        ws = wb[sheet_name]
        parts.append(f"Sheet: {sheet_name}")
        # First non-empty row = headers
        for row in ws.iter_rows(max_row=5, values_only=True):
            non_empty = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if len(non_empty) >= 2:
                parts.append(f"  Headers: {' | '.join(non_empty[:8])}")
                break
    wb.close()
    result = "\n".join(parts)
    return result[:500] if result else "(Template trống)"


def _extract_xls_structure(file_bytes: bytes) -> str:
    """Extract structure từ legacy .xls dùng pandas + xlrd."""
    import pandas as pd

    try:
        # xlrd chỉ support .xls (BIFF), không support .xlsx
        xls = pd.ExcelFile(io.BytesIO(file_bytes), engine="xlrd")
    except Exception as e:
        return f"(Không đọc được .xls: {str(e)[:80]})"

    parts = []
    for sheet_name in xls.sheet_names[:3]:
        try:
            df = pd.read_excel(xls, sheet_name=sheet_name, nrows=5, engine="xlrd")
            parts.append(f"Sheet: {sheet_name}")
            cols = [str(c).strip() for c in df.columns if not str(c).startswith("Unnamed")]
            if cols:
                parts.append(f"  Columns: {' | '.join(cols[:8])}")
        except Exception:
            pass
    result = "\n".join(parts)
    return result[:500] if result else "(Template trống)"


# ============================================================
# PROCESS NAME LOOKUP
# ============================================================

def build_process_name_lookup(inventory_rows: list[dict], sop_md_dir: Path) -> dict:
    """
    Build map: process_code → process_name.
    Source: 
      1. SOP markdown files đã OCR (parse H1)
      2. Folder name pattern (fallback)
    """
    lookup = {}

    # 1. Parse SOP markdown files
    if sop_md_dir.exists():
        for md_file in sop_md_dir.glob("*.md"):
            try:
                content = md_file.read_text(encoding="utf-8")
                # Parse H1 (tên quy trình) + Mã quy trình line
                h1_match = re.search(r"^# (.+?)$", content, re.MULTILINE)
                code_match = re.search(r"\*\*Mã quy trình:\*\*\s*(\S+)", content)
                if h1_match and code_match:
                    code = code_match.group(1).strip()
                    if code != "(không":  # filter "(không có trong lưu đồ)"
                        lookup[code.upper()] = h1_match.group(1).strip()
            except Exception as e:
                logger.debug(f"Failed to parse {md_file}: {e}")

    # 2. Fallback từ folder name
    for row in inventory_rows:
        pc = row.get("process_code", "").upper()
        if pc and pc not in lookup:
            folder = row.get("parent_folder_name", "")
            # Folder pattern: "11. QT tuyen dung" → "Quy trình tuyển dụng"
            m = re.match(r"^\d+\.\s*(.+)", folder.strip())
            if m:
                name = m.group(1).strip()
                # Expand "QT" → "Quy trình"
                name = re.sub(r"^QT\s+", "Quy trình ", name, flags=re.IGNORECASE)
                lookup[pc] = name

    return lookup


# ============================================================
# LLM CALL
# ============================================================

def _call_openai_llm(prompt_filled: str, config: dict) -> str:
    """Gọi OpenAI LLM sinh form card."""
    from openai import OpenAI

    client = OpenAI(api_key=config["llm"]["api_key"])
    response = client.chat.completions.create(
        model=config["llm"]["model"],
        messages=[{"role": "user", "content": prompt_filled}],
        temperature=0.2,
        max_tokens=4000,
    )
    return response.choices[0].message.content


def _call_gemini_llm(prompt_filled: str, config: dict) -> str:
    """Gọi Gemini LLM sinh form card."""
    import google.generativeai as genai

    genai.configure(api_key=config["llm"]["api_key"])
    model = genai.GenerativeModel(config["llm"]["model"])
    response = model.generate_content(
        prompt_filled,
        generation_config={
            "temperature": 0.2,
            "max_output_tokens": 4000,
        },
    )
    return response.text


def _call_claude_cli_llm(prompt_filled: str, config: dict) -> str:
    """Gọi Claude Code CLI (claude -p) sinh form card."""
    model = config["llm"].get("model", "claude-sonnet-4-6")
    max_turns = config["llm"].get("max_turns", 1)
    timeout = config["llm"].get("cli_timeout_seconds", 120)

    cmd = [
        "claude",
        "-p", prompt_filled,
        "--output-format", "text",
        "--model", model,
        "--max-turns", str(max_turns),
    ]

    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=timeout,
    )

    if result.returncode != 0:
        stderr = result.stderr.strip()[:500]
        raise RuntimeError(f"Claude CLI exited with code {result.returncode}: {stderr}")

    output = result.stdout.strip()
    if not output:
        raise RuntimeError("Claude CLI returned empty output")

    return output


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=2, min=2, max=20),
    reraise=True,
)
def call_llm_form_card(prompt_filled: str, config: dict) -> str:
    """Dispatch LLM theo provider, có retry."""
    provider = config["llm"]["provider"].lower()
    if provider == "openai":
        return _call_openai_llm(prompt_filled, config)
    elif provider == "gemini":
        return _call_gemini_llm(prompt_filled, config)
    elif provider == "claude":
        return _call_claude_cli_llm(prompt_filled, config)
    else:
        raise ValueError(f"Unknown LLM provider: {provider}")


# ============================================================
# VALIDATION
# ============================================================

REQUIRED_FIELDS = ["form_code", "download_url", "department"]


def validate_required_fields(file_record: dict) -> tuple[bool, list[str]]:
    """Check 3 field bắt buộc. Returns (is_valid, missing_fields)."""
    missing = []
    if not file_record.get("form_code"):
        missing.append("form_code")
    if not file_record.get("web_view_link"):
        missing.append("download_url")
    if not file_record.get("department") or file_record["department"] == "Không xác định":
        missing.append("department")
    return (len(missing) == 0, missing)


# ============================================================
# OUTPUT POSTPROCESSING
# ============================================================

def append_source_metadata(markdown: str, file_record: dict, process_name: str) -> str:
    """Append metadata footer cho RAGFlow ingest."""
    footer = (
        f"\n\n---\n\n"
        f"<!-- Source metadata (do not edit) -->\n"
        f"<!-- form_code: {file_record['form_code']} -->\n"
        f"<!-- department: {file_record['department']} -->\n"
        f"<!-- process_code: {file_record['process_code']} -->\n"
        f"<!-- process_name: {process_name} -->\n"
        f"<!-- download_url: {file_record['web_view_link']} -->\n"
        f"<!-- file_format: {file_record['file_extension']} -->\n"
        f"<!-- file_size_kb: {file_record['file_size_kb']} -->\n"
        f"<!-- last_modified: {file_record['modified_time']} -->\n"
        f"<!-- source_file_id: {file_record['file_id']} -->\n"
    )
    return markdown.rstrip() + footer


# ============================================================
# MAIN
# ============================================================

def load_form_files_from_inventory(inventory_path: Path) -> list[dict]:
    rows = []
    with open(inventory_path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row["target_kb"] == "forms_kb" and row["is_duplicate"] != "yes":
                rows.append(row)
    return rows


def process_one_form(
    file_record: dict,
    service,
    prompt_template: str,
    process_lookup: dict,
    config: dict,
    output_dir: Path,
) -> dict:
    """Xử lý 1 biểu mẫu → form card markdown."""
    log = {
        "file_id": file_record["file_id"],
        "file_name": file_record["file_name"],
        "form_code": file_record["form_code"],
        "department": file_record["department"],
        "status": "",
        "output_path": "",
        "error": "",
    }

    # 1. Validate required fields
    is_valid, missing = validate_required_fields(file_record)
    if not is_valid:
        log["status"] = "rejected_missing_fields"
        log["error"] = f"Missing: {','.join(missing)}"
        return log

    try:
        # 2. Download + extract structure
        file_bytes = gdrive_download_bytes(service, file_record["file_id"])
        ext = file_record["file_extension"].lower()
        structure = extract_doc_structure(file_bytes, ext)

        # 3. Lookup process name
        process_code = file_record["process_code"] or "(không xác định)"
        process_name = process_lookup.get(process_code.upper(), "(chưa xác định)")

        # 4. Fill prompt template
        prompt_filled = (
            prompt_template
            .replace("{form_code}", file_record["form_code"])
            .replace("{file_name}", file_record["file_name"])
            .replace("{department}", file_record["department"])
            .replace("{process_code}", process_code)
            .replace("{process_name}", process_name)
            .replace("{file_format}", ext)
            .replace("{form_structure}", structure)
            .replace("{download_url}", file_record["web_view_link"])
        )

        # 5. Call LLM
        markdown = call_llm_form_card(prompt_filled, config)

        # 6. Validate output có chứa required fields trong content
        if file_record["form_code"] not in markdown:
            log["status"] = "warning_missing_form_code"
            logger.warning(f"⚠ {file_record['form_code']}: form_code không xuất hiện trong output")
        elif file_record["web_view_link"] not in markdown:
            log["status"] = "warning_missing_url"
            logger.warning(f"⚠ {file_record['form_code']}: download_url không xuất hiện trong output")
        else:
            log["status"] = "success"

        # 7. Append metadata footer
        markdown = append_source_metadata(markdown, file_record, process_name)

        # 8. Write
        out_filename = safe_filename(f"{file_record['form_code']}.md")
        out_path = output_dir / out_filename
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(markdown)
        log["output_path"] = str(out_path.relative_to(PROJECT_ROOT))

    except Exception as e:
        log["status"] = "error"
        log["error"] = str(e)[:500]
        logger.error(f"✗ {file_record['form_code']}: {e}")

    return log


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=None)
    parser.add_argument("--resume", action="store_true")
    parser.add_argument("--config", default=None)
    args = parser.parse_args()

    config = load_config(args.config)
    inventory_csv = Path(config["paths"]["inventory_csv"])
    if not inventory_csv.is_absolute():
        inventory_csv = PROJECT_ROOT / inventory_csv
    forms_md_dir = Path(config["paths"]["forms_md_dir"])
    if not forms_md_dir.is_absolute():
        forms_md_dir = PROJECT_ROOT / forms_md_dir
    forms_md_dir.mkdir(parents=True, exist_ok=True)
    sop_md_dir = Path(config["paths"]["sop_md_dir"])
    if not sop_md_dir.is_absolute():
        sop_md_dir = PROJECT_ROOT / sop_md_dir

    generation_log = forms_md_dir.parent / "forms_generation_log.csv"
    rejected_log = forms_md_dir.parent / "forms_kb_rejected.csv"

    # Load prompt
    prompt_path = PROJECT_ROOT / "prompts" / "form_metadata_gen.txt"
    prompt_template = prompt_path.read_text(encoding="utf-8")

    # Load inventory
    if not inventory_csv.exists():
        logger.error(f"Inventory không tồn tại: {inventory_csv}")
        sys.exit(1)
    files = load_form_files_from_inventory(inventory_csv)
    logger.info(f"Tìm thấy {len(files)} biểu mẫu cho forms_kb")

    if args.resume:
        before = len(files)
        files = [
            f for f in files
            if f["form_code"]
            and not (forms_md_dir / safe_filename(f"{f['form_code']}.md")).exists()
        ]
        logger.info(f"Resume: skip {before - len(files)} file đã có .md")

    if args.limit:
        files = files[: args.limit]
        logger.info(f"Limit: {len(files)} file đầu")

    if not files:
        logger.info("Không có file nào để xử lý.")
        return

    # Build process lookup
    logger.info("Building process_code → name lookup...")
    all_form_rows = load_form_files_from_inventory(inventory_csv)
    process_lookup = build_process_name_lookup(all_form_rows, sop_md_dir)
    logger.info(f"Lookup table: {len(process_lookup)} process codes")

    # Cost estimate
    llm_provider = config["llm"]["provider"].lower()
    if llm_provider == "openai":
        per_file = 0.010 if "mini" in config["llm"]["model"] else 0.03
        est_cost = len(files) * per_file
        logger.info(f"💰 Ước tính chi phí OpenAI ({config['llm']['model']}): ~${est_cost:.2f}")
    elif llm_provider == "gemini":
        est_cost = len(files) * 0.002
        logger.info(f"💰 Ước tính chi phí Gemini: ~${est_cost:.3f}")
    elif llm_provider == "claude":
        est_cost = len(files) * 0.015
        logger.info(f"💰 Ước tính chi phí Claude CLI: ~${est_cost:.2f}")

    # Process
    service = get_gdrive_client(config)
    rate_delay = 1.0 / config["llm"]["rate_limit_rps"]
    logs = []
    rejected = []

    for f in tqdm(files, desc="Generate form cards"):
        log = process_one_form(f, service, prompt_template, process_lookup, config, forms_md_dir)
        if log["status"] == "rejected_missing_fields":
            rejected.append({**f, "missing": log["error"]})
        logs.append(log)
        time.sleep(rate_delay)

    # Write generation log
    fieldnames_log = ["file_id", "file_name", "form_code", "department", "status", "output_path", "error"]
    file_exists = generation_log.exists() and args.resume
    with open(generation_log, "a" if file_exists else "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=fieldnames_log)
        if not file_exists:
            w.writeheader()
        w.writerows(logs)

    # Write rejected log
    if rejected:
        with open(rejected_log, "w", newline="", encoding="utf-8") as fh:
            w = csv.DictWriter(fh, fieldnames=list(rejected[0].keys()))
            w.writeheader()
            w.writerows(rejected)
        logger.warning(f"⚠ {len(rejected)} file rejected → {rejected_log}")

    # Summary
    by_status = {}
    for l in logs:
        by_status[l["status"]] = by_status.get(l["status"], 0) + 1
    logger.info("=" * 60)
    logger.info("FORM CARDS GENERATION SUMMARY")
    logger.info("=" * 60)
    for s, c in sorted(by_status.items()):
        logger.info(f"  {s}: {c}")
    logger.info(f"Output dir: {forms_md_dir}")
    logger.info(f"Generation log: {generation_log}")
    if rejected:
        logger.info(f"Rejected files: {rejected_log} ({len(rejected)} files)")
    logger.info("=" * 60)


if __name__ == "__main__":
    main()
